"""
Tests for PDFScanner
=====================
Run from the project root with:
    pytest tests/
    pytest tests/ -v          # verbose
    pytest tests/ -v -k ocr  # run only tests whose name contains "ocr"

Most tests use mocks so they run without PaddleOCR, OpenAI, or real files.
The integration tests at the bottom are skipped unless you pass --run-integration.
"""

import io
import json
import os
import tempfile
import textwrap
import types
import urllib.parse
from pathlib import Path
from unittest.mock import MagicMock, mock_open, patch

import pytest


# ---------------------------------------------------------------------------
# Helpers / fixtures
# ---------------------------------------------------------------------------

SAMPLES_DIR = Path(__file__).resolve().parent.parent / "data" / "samples"


# The paddleocr module is stubbed in conftest.py before any test imports
# agent.text_extraction, so unit tests never load the real model.

def _fake_ocr_result(texts: list[str]):
    """Return a list of fake PaddleOCR result objects matching the 3.x API."""
    results = []
    for text in texts:
        res = MagicMock()
        res.json = {"res": {"rec_texts": [text]}}
        results.append(res)
    return results


# ---------------------------------------------------------------------------
# 1. paddle_predict_to_text
# ---------------------------------------------------------------------------

class TestPaddlePredictToText:
    """Unit tests for the PaddleOCR output parser."""

    def _import(self):
        from agent.text_extraction import paddle_predict_to_text
        return paddle_predict_to_text

    def test_normal_rec_texts(self):
        fn = self._import()
        res = MagicMock()
        res.json = {"res": {"rec_texts": ["hello", "world"]}}
        assert fn([res]) == "hello\nworld"

    def test_empty_input(self):
        fn = self._import()
        assert fn([]) == ""

    def test_result_without_json_attr(self):
        fn = self._import()
        res = MagicMock(spec=[])          # no .json attribute
        assert fn([res]) == ""

    def test_json_is_not_a_dict(self):
        fn = self._import()
        res = MagicMock()
        res.json = "not a dict"
        assert fn([res]) == ""

    def test_nested_list_of_dicts(self):
        """Covers the multi-page / multi-stage branch."""
        fn = self._import()
        res = MagicMock()
        res.json = {
            "res": [
                {"rec_texts": ["page one line"]},
                {"rec_texts": ["page two line"]},
            ]
        }
        result = fn([res])
        assert "page one line" in result
        assert "page two line" in result

    def test_skips_non_string_rec_texts(self):
        fn = self._import()
        res = MagicMock()
        res.json = {"res": {"rec_texts": ["good", None, 42, "also good"]}}
        result = fn([res])
        assert result == "good\nalso good"

    def test_res_key_falls_back_to_root(self):
        """If there's no 'res' key, it should use the top-level dict."""
        fn = self._import()
        res = MagicMock()
        res.json = {"rec_texts": ["fallback text"]}
        assert fn([res]) == "fallback text"


# ---------------------------------------------------------------------------
# 2. extract_text — routing logic
# ---------------------------------------------------------------------------

class TestExtractTextRouting:
    """Checks that extract_text calls the right sub-function per extension."""

    def _get_module(self, monkeypatch):
        import agent.text_extraction as te
        monkeypatch.setattr(te, "ocr", MagicMock())
        return te

    def test_pdf_routes_to_extract_pdf(self, monkeypatch):
        te = self._get_module(monkeypatch)
        monkeypatch.setattr(te, "extract_pdf", MagicMock(return_value="pdf text long enough"))
        result = te.extract_text("doc.pdf")
        te.extract_pdf.assert_called_once_with("doc.pdf")
        assert result == "pdf text long enough"

    def test_png_routes_to_ocr_image(self, monkeypatch):
        te = self._get_module(monkeypatch)
        monkeypatch.setattr(te, "ocr_image", MagicMock(return_value="ocr result"))
        result = te.extract_text("scan.png")
        te.ocr_image.assert_called_once_with("scan.png")
        assert result == "ocr result"

    def test_jpg_routes_to_ocr_image(self, monkeypatch):
        te = self._get_module(monkeypatch)
        monkeypatch.setattr(te, "ocr_image", MagicMock(return_value="ocr result"))
        te.extract_text("photo.jpg")
        te.ocr_image.assert_called_once()

    def test_jpeg_routes_to_ocr_image(self, monkeypatch):
        te = self._get_module(monkeypatch)
        monkeypatch.setattr(te, "ocr_image", MagicMock(return_value="ocr result"))
        te.extract_text("photo.jpeg")
        te.ocr_image.assert_called_once()

    def test_docx_routes_to_extract_docx(self, monkeypatch):
        te = self._get_module(monkeypatch)
        monkeypatch.setattr(te, "extract_docx", MagicMock(return_value="docx text"))
        result = te.extract_text("report.docx")
        te.extract_docx.assert_called_once_with("report.docx")
        assert result == "docx text"

    def test_unsupported_extension_raises(self, monkeypatch):
        te = self._get_module(monkeypatch)
        with pytest.raises(ValueError, match="Unsupported file type"):
            te.extract_text("data.xlsx")

# ---------------------------------------------------------------------------
# 2b. extract_pdf
# ---------------------------------------------------------------------------

class TestExtractPdf:
    def _fake_pdf(self, page_texts):
        pages = []
        for text in page_texts:
            page = MagicMock()
            page.extract_text.return_value = text
            pages.append(page)
        fake_pdf = MagicMock()
        fake_pdf.pages = pages
        return fake_pdf

    def test_pages_are_separated_by_blank_line(self):
        from agent.text_extraction import extract_pdf

        with (
            patch("agent.text_extraction.pdfplumber.open") as mock_open_pdf,
            patch("agent.text_extraction.ocr_pdf_pages") as mock_ocr,
        ):
            mock_open_pdf.return_value.__enter__.return_value = self._fake_pdf(
                ["Page one has enough embedded text", "Page two also has enough text"]
            )
            assert extract_pdf("doc.pdf") == (
                "Page one has enough embedded text\n\nPage two also has enough text"
            )
            mock_ocr.assert_not_called()

    def test_mixed_pdf_ocrs_only_pages_with_little_text(self):
        from agent.text_extraction import extract_pdf

        with (
            patch("agent.text_extraction.pdfplumber.open") as mock_open_pdf,
            patch(
                "agent.text_extraction.ocr_pdf_pages",
                return_value={1: "OCR text from scanned page"},
            ) as mock_ocr,
        ):
            mock_open_pdf.return_value.__enter__.return_value = self._fake_pdf(
                ["Digital cover page with selectable text", "watermark", "Another digital page"]
            )
            assert extract_pdf("doc.pdf") == (
                "Digital cover page with selectable text\n\n"
                "OCR text from scanned page\n\n"
                "Another digital page"
            )
            mock_ocr.assert_called_once_with("doc.pdf", [1])

    def test_keeps_short_embedded_text_when_ocr_returns_empty(self):
        from agent.text_extraction import extract_pdf

        with (
            patch("agent.text_extraction.pdfplumber.open") as mock_open_pdf,
            patch("agent.text_extraction.ocr_pdf_pages", return_value={0: ""}),
        ):
            mock_open_pdf.return_value.__enter__.return_value = self._fake_pdf(["Short title"])
            assert extract_pdf("doc.pdf") == "Short title"


# ---------------------------------------------------------------------------
# 3. ocr_pdf fallback
# ---------------------------------------------------------------------------

class TestOcrPdfFallback:
    def test_rendered_pages_use_unique_directory_and_are_cleaned_up(self, monkeypatch):
        import sys

        import agent.text_extraction as te

        rendered_paths = []

        class FakePixmap:
            def save(self, path):
                Path(path).write_bytes(b"rendered page")

        class FakePage:
            def get_pixmap(self):
                return FakePixmap()

        class FakePdf:
            def __enter__(self):
                return self

            def __exit__(self, exc_type, exc, traceback):
                return False

            def __len__(self):
                return 2

            def __getitem__(self, index):
                return FakePage()

        fake_fitz = types.ModuleType("fitz")
        fake_fitz.open = MagicMock(return_value=FakePdf())
        monkeypatch.setitem(sys.modules, "fitz", fake_fitz)

        fake_ocr = MagicMock()

        def predict(*, input):
            rendered_path = Path(input)
            assert rendered_path.exists()
            rendered_paths.append(rendered_path)
            page_number = int(rendered_path.stem.split("_")[-1]) + 1
            return _fake_ocr_result([f"page {page_number}"])

        fake_ocr.predict.side_effect = predict
        monkeypatch.setattr(te, "ocr", fake_ocr)

        first_result = te.ocr_pdf("scan.pdf")
        second_result = te.ocr_pdf("scan.pdf")

        assert first_result == "page 1\n\npage 2"
        assert second_result == "page 1\n\npage 2"
        temp_dirs = {path.parent for path in rendered_paths}
        assert len(temp_dirs) == 2
        assert all(temp_dir != Path(tempfile.gettempdir()) for temp_dir in temp_dirs)
        assert all(not temp_dir.exists() for temp_dir in temp_dirs)


# ---------------------------------------------------------------------------
# 4. extract_docx
# ---------------------------------------------------------------------------

class TestExtractDocx:
    """Uses real .docx files built with python-docx instead of mocks, so the
    tests exercise the actual table/header/footer traversal."""

    def _write_docx(self, path, build):
        import docx as docx_lib

        document = docx_lib.Document()
        build(document)
        document.save(str(path))

    def test_joins_paragraphs(self, tmp_path):
        from agent.text_extraction import extract_docx

        docx_path = tmp_path / "paras.docx"
        self._write_docx(
            docx_path,
            lambda d: (d.add_paragraph("Hello"), d.add_paragraph("World")),
        )

        assert extract_docx(str(docx_path)) == "Hello\nWorld"

    def test_empty_document(self, tmp_path):
        from agent.text_extraction import extract_docx

        docx_path = tmp_path / "empty.docx"
        self._write_docx(docx_path, lambda d: None)

        assert extract_docx(str(docx_path)).strip() == ""

    def test_includes_table_content_with_cells_joined_per_row(self, tmp_path):
        from agent.text_extraction import extract_docx

        def build(document):
            document.add_paragraph("Invoice #42")
            table = document.add_table(rows=2, cols=2)
            table.rows[0].cells[0].text = "Consulting"
            table.rows[0].cells[1].text = "$1,200.00"
            table.rows[1].cells[0].text = "Total"
            table.rows[1].cells[1].text = "$1,450.00"

        docx_path = tmp_path / "invoice.docx"
        self._write_docx(docx_path, build)

        result = extract_docx(str(docx_path))
        assert "Consulting | $1,200.00" in result
        assert "Total | $1,450.00" in result

    def test_preserves_document_order_around_tables(self, tmp_path):
        from agent.text_extraction import extract_docx

        def build(document):
            document.add_paragraph("Before the table")
            table = document.add_table(rows=1, cols=1)
            table.rows[0].cells[0].text = "Inside the table"
            document.add_paragraph("After the table")

        docx_path = tmp_path / "ordered.docx"
        self._write_docx(docx_path, build)

        result = extract_docx(str(docx_path))
        assert (
            result.index("Before the table")
            < result.index("Inside the table")
            < result.index("After the table")
        )

    def test_includes_header_and_footer_text(self, tmp_path):
        from agent.text_extraction import extract_docx

        def build(document):
            section = document.sections[0]
            section.header.paragraphs[0].text = "Acme Corp - 123 Main St"
            section.footer.paragraphs[0].text = "Page footer notice"
            document.add_paragraph("Body text")

        docx_path = tmp_path / "letterhead.docx"
        self._write_docx(docx_path, build)

        result = extract_docx(str(docx_path))
        assert "Acme Corp - 123 Main St" in result
        assert "Body text" in result
        assert "Page footer notice" in result


# ---------------------------------------------------------------------------
# 4. classify_document
# ---------------------------------------------------------------------------

class TestClassifyDocument:
    def test_returns_parsed_classification(self, monkeypatch):
        from agent.doc_classify import DocumentClassification, classify_document

        fake_response = MagicMock()
        fake_response.output_parsed = DocumentClassification(
            document_type="invoice",
            confidence=95,
        )

        fake_client = MagicMock()
        fake_client.responses.parse.return_value = fake_response

        monkeypatch.setattr("agent.doc_classify._client", lambda: fake_client)

        result = classify_document("Invoice #1234 for $500")
        assert result == {"document_type": "invoice", "confidence": 95}
        call_kwargs = fake_client.responses.parse.call_args.kwargs
        assert call_kwargs["text_format"] is DocumentClassification

    def test_truncates_text_to_3000_chars(self, monkeypatch):
        from agent.doc_classify import DocumentClassification, classify_document

        fake_response = MagicMock()
        fake_response.output_parsed = DocumentClassification(
            document_type="other",
            confidence=50,
        )

        fake_client = MagicMock()
        fake_client.responses.parse.return_value = fake_response
        monkeypatch.setattr("agent.doc_classify._client", lambda: fake_client)

        long_text = "x" * 10_000
        classify_document(long_text)

        call_args = fake_client.responses.parse.call_args
        assert call_args.kwargs["input"] == long_text[:3000]

    def test_raises_without_parsed_classification(self, monkeypatch):
        from agent.doc_classify import classify_document

        fake_response = MagicMock()
        fake_response.output_parsed = None
        fake_client = MagicMock()
        fake_client.responses.parse.return_value = fake_response
        monkeypatch.setattr("agent.doc_classify._client", lambda: fake_client)

        with pytest.raises(ValueError, match="classification model"):
            classify_document("some document text")


# ---------------------------------------------------------------------------
# 5. summarize_document
# ---------------------------------------------------------------------------

class TestSummarizeDocument:
    def test_returns_parsed_summary(self, monkeypatch):
        from agent.doc_summarize import DocumentSummary, summarize_document

        fake_response = MagicMock()
        fake_response.output_parsed = DocumentSummary(
            short_summary="A test document.",
            bullet_points=["Point A", "Point B"],
        )

        fake_client = MagicMock()
        fake_client.responses.parse.return_value = fake_response
        monkeypatch.setattr("agent.doc_summarize._client", lambda: fake_client)

        result = summarize_document("Some document text")
        assert result == {
            "short_summary": "A test document.",
            "bullet_points": ["Point A", "Point B"],
        }
        call_kwargs = fake_client.responses.parse.call_args.kwargs
        assert call_kwargs["text_format"] is DocumentSummary

    def test_truncates_text_to_4000_chars(self, monkeypatch):
        from agent.doc_summarize import DocumentSummary, summarize_document

        fake_response = MagicMock()
        fake_response.output_parsed = DocumentSummary(
            short_summary="Short.",
            bullet_points=[],
        )
        fake_client = MagicMock()
        fake_client.responses.parse.return_value = fake_response
        monkeypatch.setattr("agent.doc_summarize._client", lambda: fake_client)

        long_text = "y" * 10_000
        summarize_document(long_text)

        call_args = fake_client.responses.parse.call_args
        assert call_args.kwargs["input"] == long_text[:4000]

    def test_raises_without_parsed_summary(self, monkeypatch):
        from agent.doc_summarize import summarize_document

        fake_response = MagicMock()
        fake_response.output_parsed = None
        fake_client = MagicMock()
        fake_client.responses.parse.return_value = fake_response
        monkeypatch.setattr("agent.doc_summarize._client", lambda: fake_client)

        with pytest.raises(ValueError, match="summary model"):
            summarize_document("some document text")


# ---------------------------------------------------------------------------
# 5b. doc_explain
# ---------------------------------------------------------------------------

class TestDocExplain:
    def _explanation(self):
        from agent.doc_explain import DocumentExplanation
        return DocumentExplanation(
            document_type="Invoice",
            summary="An invoice for services.",
            explanation="The document requests payment.",
            amounts=["$500"],
        )

    def _fake_client(self, parsed):
        response = MagicMock()
        response.output_parsed = parsed
        fake_client = MagicMock()
        fake_client.responses.parse.return_value = response
        return fake_client

    def test_explain_document_returns_dict(self, monkeypatch):
        from agent import doc_explain

        fake_client = self._fake_client(self._explanation())
        monkeypatch.setattr(doc_explain, "_client", lambda: fake_client)

        result = doc_explain.explain_document("some document text")

        assert result["document_type"] == "Invoice"
        assert result["amounts"] == ["$500"]
        assert result["warnings"] == []
        call_kwargs = fake_client.responses.parse.call_args.kwargs
        assert call_kwargs["input"] == "some document text"
        assert "blank or mostly empty form" in call_kwargs["instructions"]

    def test_explain_document_raises_without_parsed_output(self, monkeypatch):
        from agent import doc_explain

        fake_client = self._fake_client(None)
        monkeypatch.setattr(doc_explain, "_client", lambda: fake_client)

        with pytest.raises(ValueError):
            doc_explain.explain_document("some text")

    def test_translate_explanation_targets_language(self, monkeypatch):
        from agent import doc_explain

        fake_client = self._fake_client(self._explanation())
        monkeypatch.setattr(doc_explain, "_client", lambda: fake_client)

        original = self._explanation().model_dump()
        result = doc_explain.translate_explanation(original, "Ukrainian")

        assert result["document_type"] == "Invoice"
        call_kwargs = fake_client.responses.parse.call_args.kwargs
        assert "Ukrainian" in call_kwargs["instructions"]
        assert json.loads(call_kwargs["input"]) == original

    def test_translate_explanation_raises_without_parsed_output(self, monkeypatch):
        from agent import doc_explain

        fake_client = self._fake_client(None)
        monkeypatch.setattr(doc_explain, "_client", lambda: fake_client)

        with pytest.raises(ValueError):
            doc_explain.translate_explanation({"summary": "x"}, "French")


# ---------------------------------------------------------------------------
# 6. Flask web app
# ---------------------------------------------------------------------------

@pytest.fixture()
def flask_client(monkeypatch, tmp_path):
    """
    Create a Flask test client with extract_text mocked so no real OCR runs.
    UPLOAD_DIR is redirected to tmp_path.
    """
    import sys
    fake = types.ModuleType("paddleocr")
    fake.PaddleOCR = MagicMock(return_value=MagicMock())
    sys.modules.setdefault("paddleocr", fake)

    import web_app
    monkeypatch.setattr("web_app.UPLOAD_DIR", tmp_path / "uploads")
    monkeypatch.setattr("web_app.extract_text", MagicMock(return_value="extracted text content"))
    monkeypatch.setattr("web_app.validate_pdf_page_count", MagicMock())
    monkeypatch.setattr(web_app.limiter, "enabled", False)
    monkeypatch.setattr("web_app.TURNSTILE_ENABLED", False)
    monkeypatch.setattr("web_app.OPENAI_FEATURES_ENABLED", True)
    web_app.limiter.reset()

    web_app.app.config["TESTING"] = True
    with web_app.app.test_client() as client:
        yield client, tmp_path


class TestWebUploadCleanup:
    def test_cleanup_upload_dir_removes_stale_upload_files(self, monkeypatch, tmp_path):
        import web_app

        upload_dir = tmp_path / "uploads"
        upload_dir.mkdir()
        stale_file = upload_dir / "stale.pdf"
        stale_file.write_bytes(b"leftover upload")
        nested_dir = upload_dir / "nested"
        nested_dir.mkdir()
        nested_file = nested_dir / "keep.txt"
        nested_file.write_text("not an upload file")
        monkeypatch.setattr("web_app.UPLOAD_DIR", upload_dir)

        web_app.cleanup_upload_dir()

        assert not stale_file.exists()
        assert nested_file.exists()

    def test_cleanup_upload_dir_allows_missing_directory(self, monkeypatch, tmp_path):
        import web_app

        monkeypatch.setattr("web_app.UPLOAD_DIR", tmp_path / "missing-uploads")

        web_app.cleanup_upload_dir()


class TestWebPdfPageValidation:
    def test_validate_pdf_page_count_allows_pdf_at_limit(self, monkeypatch, tmp_path):
        import sys
        import web_app

        class FakePdf:
            page_count = web_app.MAX_PDF_PAGES

            def __enter__(self):
                return self

            def __exit__(self, exc_type, exc, traceback):
                return False

        fake_fitz = types.ModuleType("fitz")
        fake_fitz.FileDataError = RuntimeError
        fake_fitz.open = MagicMock(return_value=FakePdf())
        monkeypatch.setitem(sys.modules, "fitz", fake_fitz)

        pdf_path = tmp_path / "doc.pdf"
        pdf_path.write_bytes(b"dummy")

        web_app.validate_pdf_page_count(pdf_path)
        fake_fitz.open.assert_called_once_with(pdf_path)

    def test_validate_pdf_page_count_rejects_pdf_over_limit(self, monkeypatch, tmp_path):
        import sys
        import web_app

        class FakePdf:
            page_count = web_app.MAX_PDF_PAGES + 1

            def __enter__(self):
                return self

            def __exit__(self, exc_type, exc, traceback):
                return False

        fake_fitz = types.ModuleType("fitz")
        fake_fitz.FileDataError = RuntimeError
        fake_fitz.open = MagicMock(return_value=FakePdf())
        monkeypatch.setitem(sys.modules, "fitz", fake_fitz)

        pdf_path = tmp_path / "doc.pdf"
        pdf_path.write_bytes(b"dummy")

        with pytest.raises(web_app.PdfPageLimitError, match="limited"):
            web_app.validate_pdf_page_count(pdf_path)

    def test_validate_pdf_page_count_rejects_unreadable_pdf(self, monkeypatch, tmp_path):
        import sys
        import web_app

        class FakeFileDataError(Exception):
            pass

        fake_fitz = types.ModuleType("fitz")
        fake_fitz.FileDataError = FakeFileDataError
        fake_fitz.open = MagicMock(side_effect=FakeFileDataError("bad pdf"))
        monkeypatch.setitem(sys.modules, "fitz", fake_fitz)

        pdf_path = tmp_path / "corrupt.pdf"
        pdf_path.write_bytes(b"not a pdf")

        with pytest.raises(web_app.PdfValidationError, match="corrupt or unreadable"):
            web_app.validate_pdf_page_count(pdf_path)


class TestWebUploadResourceValidation:
    def test_validate_image_allows_image_at_pixel_limit(self, monkeypatch, tmp_path):
        import web_app

        class FakeImage:
            size = (web_app.MAX_IMAGE_PIXELS, 1)

            def __enter__(self):
                return self

            def __exit__(self, exc_type, exc, traceback):
                return False

        mock_open = MagicMock(return_value=FakeImage())
        monkeypatch.setattr("PIL.Image.open", mock_open)

        image_path = tmp_path / "scan.png"
        image_path.write_bytes(b"not decoded by test")

        web_app.validate_image(image_path)
        mock_open.assert_called_once_with(image_path)

    def test_validate_image_rejects_image_over_pixel_limit(self, monkeypatch, tmp_path):
        import web_app

        class FakeImage:
            size = (web_app.MAX_IMAGE_PIXELS + 1, 1)

            def __enter__(self):
                return self

            def __exit__(self, exc_type, exc, traceback):
                return False

        monkeypatch.setattr("PIL.Image.open", MagicMock(return_value=FakeImage()))

        image_path = tmp_path / "huge.jpg"
        image_path.write_bytes(b"not decoded by test")

        with pytest.raises(web_app.PdfValidationError, match="dimensions are too large"):
            web_app.validate_image(image_path)

    def test_validate_image_rejects_corrupt_image(self, monkeypatch, tmp_path):
        from PIL import UnidentifiedImageError
        import web_app

        monkeypatch.setattr(
            "PIL.Image.open",
            MagicMock(side_effect=UnidentifiedImageError("bad image")),
        )

        image_path = tmp_path / "corrupt.png"
        image_path.write_bytes(b"not an image")

        with pytest.raises(web_app.PdfValidationError, match="corrupt or not a real image"):
            web_app.validate_image(image_path)

    def test_validate_docx_allows_docx_at_uncompressed_limit(self, monkeypatch, tmp_path):
        import web_app

        fake_info = MagicMock(file_size=web_app.MAX_DOCX_UNCOMPRESSED)
        fake_zip = MagicMock()
        fake_zip.__enter__.return_value.infolist.return_value = [fake_info]
        fake_zip.__exit__.return_value = False
        mock_zip = MagicMock(return_value=fake_zip)
        monkeypatch.setattr("web_app.zipfile.ZipFile", mock_zip)

        docx_path = tmp_path / "doc.docx"
        docx_path.write_bytes(b"not opened by test")

        web_app.validate_docx(docx_path)
        mock_zip.assert_called_once_with(docx_path)

    def test_validate_docx_rejects_docx_over_uncompressed_limit(self, monkeypatch, tmp_path):
        import web_app

        fake_info = MagicMock(file_size=web_app.MAX_DOCX_UNCOMPRESSED + 1)
        fake_zip = MagicMock()
        fake_zip.__enter__.return_value.infolist.return_value = [fake_info]
        fake_zip.__exit__.return_value = False
        monkeypatch.setattr("web_app.zipfile.ZipFile", MagicMock(return_value=fake_zip))

        docx_path = tmp_path / "large.docx"
        docx_path.write_bytes(b"not opened by test")

        with pytest.raises(web_app.PdfValidationError, match="expands too large"):
            web_app.validate_docx(docx_path)

    def test_validate_docx_rejects_corrupt_docx(self, monkeypatch, tmp_path):
        import web_app

        monkeypatch.setattr(
            "web_app.zipfile.ZipFile",
            MagicMock(side_effect=web_app.zipfile.BadZipFile("bad zip")),
        )

        docx_path = tmp_path / "corrupt.docx"
        docx_path.write_bytes(b"not a zip")

        with pytest.raises(web_app.PdfValidationError, match="corrupt or unreadable"):
            web_app.validate_docx(docx_path)

    def test_validate_upload_resource_limits_runs_each_validator(self, monkeypatch, tmp_path):
        import web_app

        mock_pdf = MagicMock()
        mock_image = MagicMock()
        mock_docx = MagicMock()
        monkeypatch.setattr("web_app.validate_pdf_page_count", mock_pdf)
        monkeypatch.setattr("web_app.validate_image", mock_image)
        monkeypatch.setattr("web_app.validate_docx", mock_docx)

        path = tmp_path / "scan.png"
        web_app.validate_upload_resource_limits(path)

        mock_pdf.assert_called_once_with(path)
        mock_image.assert_called_once_with(path)
        mock_docx.assert_called_once_with(path)


class TestTurnstileVerification:
    def test_verify_turnstile_response_skips_when_disabled(self, monkeypatch):
        import web_app

        monkeypatch.setattr("web_app.TURNSTILE_ENABLED", False)

        result = web_app.verify_turnstile_response(None, "203.0.113.5")

        assert result == {"success": True, "skipped": True}

    def test_verify_turnstile_response_calls_siteverify(self, monkeypatch):
        import web_app

        captured = {}

        class FakeResponse:
            def __enter__(self):
                return self

            def __exit__(self, exc_type, exc, traceback):
                return False

            def read(self):
                return b'{"success": true, "hostname": "example.com"}'

        def fake_urlopen(request, timeout):
            captured["data"] = urllib.parse.parse_qs(request.data.decode("utf-8"))
            captured["timeout"] = timeout
            return FakeResponse()

        monkeypatch.setattr("web_app.TURNSTILE_ENABLED", True)
        monkeypatch.setattr("web_app.TURNSTILE_SECRET_KEY", "secret-key")
        monkeypatch.setattr("web_app.TURNSTILE_TIMEOUT_SECONDS", 7)
        monkeypatch.setattr("web_app.urllib.request.urlopen", fake_urlopen)

        result = web_app.verify_turnstile_response("client-token", "203.0.113.5")

        assert result["success"] is True
        assert captured["data"] == {
            "secret": ["secret-key"],
            "response": ["client-token"],
            "remoteip": ["203.0.113.5"],
        }
        assert captured["timeout"] == 7

    def test_verify_turnstile_response_rejects_failure(self, monkeypatch):
        import web_app

        class FakeResponse:
            def __enter__(self):
                return self

            def __exit__(self, exc_type, exc, traceback):
                return False

            def read(self):
                return b'{"success": false, "error-codes": ["invalid-input-response"]}'

        monkeypatch.setattr("web_app.TURNSTILE_ENABLED", True)
        monkeypatch.setattr("web_app.TURNSTILE_SECRET_KEY", "secret-key")
        monkeypatch.setattr("web_app.urllib.request.urlopen", MagicMock(return_value=FakeResponse()))

        with pytest.raises(web_app.TurnstileValidationError, match="Bot verification failed"):
            web_app.verify_turnstile_response("bad-token", "203.0.113.5")

    def test_verify_turnstile_response_requires_secret_when_enabled(self, monkeypatch):
        import web_app

        monkeypatch.setattr("web_app.TURNSTILE_ENABLED", True)
        monkeypatch.setattr("web_app.TURNSTILE_SECRET_KEY", "")

        with pytest.raises(web_app.TurnstileConfigError, match="not configured"):
            web_app.verify_turnstile_response("client-token", "203.0.113.5")


class TestWebApp:
    def test_security_headers_on_pages_and_api_errors(self, flask_client):
        client, _ = flask_client

        page_resp = client.get("/")
        error_resp = client.post("/api/extract", data={"mode": "extract"})

        for resp in (page_resp, error_resp):
            assert resp.headers["X-Content-Type-Options"] == "nosniff"
            assert resp.headers["X-Frame-Options"] == "DENY"
            assert resp.headers["Referrer-Policy"] == "strict-origin-when-cross-origin"
            csp = resp.headers["Content-Security-Policy"]
            assert "default-src 'self'" in csp
            assert "script-src 'self' https://challenges.cloudflare.com" in csp
            assert "frame-ancestors 'none'" in csp

    def test_healthz_returns_ok(self, flask_client):
        client, _ = flask_client
        resp = client.get("/healthz")
        assert resp.status_code == 200
        assert resp.get_json() == {"status": "ok"}

    def test_index_returns_explainer(self, flask_client):
        client, _ = flask_client
        resp = client.get("/")
        assert resp.status_code == 200
        assert b"Understand any" in resp.data
        assert b"Download PDF" in resp.data

    def test_extract_page_returns_200(self, flask_client):
        client, _ = flask_client
        resp = client.get("/extract")
        assert resp.status_code == 200
        assert b"Document" in resp.data

    def test_privacy_page_returns_200(self, flask_client):
        client, _ = flask_client
        resp = client.get("/privacy")
        assert resp.status_code == 200
        assert b"Temporary processing" in resp.data
        assert b"OpenAI" in resp.data
        assert b"does not create user accounts" in resp.data

    def test_pages_render_turnstile_when_enabled(self, flask_client, monkeypatch):
        client, _ = flask_client

        monkeypatch.setattr("web_app.TURNSTILE_ENABLED", True)
        monkeypatch.setattr("web_app.TURNSTILE_SITE_KEY", "site-key")

        index_resp = client.get("/")
        extract_resp = client.get("/extract")

        assert b"https://challenges.cloudflare.com/turnstile/v0/api.js" in index_resp.data
        assert b"cf-turnstile" in index_resp.data
        assert b"site-key" in index_resp.data
        assert b"https://challenges.cloudflare.com/turnstile/v0/api.js" in extract_resp.data
        assert b"cf-turnstile" in extract_resp.data
        assert b"site-key" in extract_resp.data

    def test_extract_no_file_returns_400(self, flask_client):
        client, _ = flask_client
        resp = client.post("/api/extract", data={"mode": "extract"})
        assert resp.status_code == 400
        assert "error" in resp.get_json()

    def test_extract_unsupported_type_returns_400(self, flask_client):
        client, _ = flask_client
        data = {
            "file": (io.BytesIO(b"data"), "file.xlsx"),
            "mode": "extract",
        }
        resp = client.post("/api/extract", data=data, content_type="multipart/form-data")
        assert resp.status_code == 400
        body = resp.get_json()
        assert "error" in body

    def test_extract_invalid_mode_returns_400(self, flask_client):
        client, _ = flask_client
        data = {
            "file": (io.BytesIO(b"%PDF-1.4"), "doc.pdf"),
            "mode": "invalid_mode",
        }
        resp = client.post("/api/extract", data=data, content_type="multipart/form-data")
        assert resp.status_code == 400

    def test_extract_rejects_oversized_image_before_ocr(self, flask_client, monkeypatch):
        client, _ = flask_client
        import web_app

        mock_extract = MagicMock(return_value="should not run")
        monkeypatch.setattr("web_app.extract_text", mock_extract)
        monkeypatch.setattr(
            "web_app.validate_image",
            MagicMock(
                side_effect=web_app.PdfValidationError(
                    "This image's dimensions are too large to process."
                )
            ),
        )

        data = {
            "file": (io.BytesIO(b"not decoded by test"), "huge.png"),
            "mode": "extract",
        }
        resp = client.post("/api/extract", data=data, content_type="multipart/form-data")

        assert resp.status_code == 400
        assert "dimensions are too large" in resp.get_json()["error"]
        mock_extract.assert_not_called()

    def test_extract_rejects_large_docx_before_extraction(self, flask_client, monkeypatch):
        client, _ = flask_client
        import web_app

        mock_extract = MagicMock(return_value="should not run")
        monkeypatch.setattr("web_app.extract_text", mock_extract)
        monkeypatch.setattr(
            "web_app.validate_docx",
            MagicMock(
                side_effect=web_app.PdfValidationError(
                    "This DOCX expands too large to process."
                )
            ),
        )

        data = {
            "file": (io.BytesIO(b"not opened by test"), "large.docx"),
            "mode": "extract",
        }
        resp = client.post("/api/extract", data=data, content_type="multipart/form-data")

        assert resp.status_code == 400
        assert "expands too large" in resp.get_json()["error"]
        mock_extract.assert_not_called()

    def test_extract_full_mode_can_be_disabled_before_upload_save(
        self, flask_client, monkeypatch
    ):
        client, tmp_path = flask_client

        monkeypatch.setattr("web_app.OPENAI_FEATURES_ENABLED", False)
        mock_extract = MagicMock(return_value="should not run")
        monkeypatch.setattr("web_app.extract_text", mock_extract)
        mock_verify = MagicMock()
        monkeypatch.setattr("web_app.verify_request_turnstile", mock_verify)

        data = {
            "file": (io.BytesIO(b"dummy"), "invoice.pdf"),
            "mode": "full",
        }
        resp = client.post("/api/extract", data=data, content_type="multipart/form-data")

        assert resp.status_code == 503
        assert "AI document features are temporarily unavailable" in resp.get_json()["error"]
        mock_verify.assert_not_called()
        mock_extract.assert_not_called()
        assert not (tmp_path / "uploads").exists()

    def test_extract_text_mode_still_works_when_openai_features_are_disabled(
        self, flask_client, monkeypatch
    ):
        client, _ = flask_client

        monkeypatch.setattr("web_app.OPENAI_FEATURES_ENABLED", False)
        mock_extract = MagicMock(return_value="plain text")
        monkeypatch.setattr("web_app.extract_text", mock_extract)

        data = {
            "file": (io.BytesIO(b"dummy"), "invoice.pdf"),
            "mode": "extract",
        }
        resp = client.post("/api/extract", data=data, content_type="multipart/form-data")

        assert resp.status_code == 200
        assert resp.get_json()["text"] == "plain text"
        mock_extract.assert_called_once()

    def test_extract_requires_turnstile_before_saving_upload(self, flask_client, monkeypatch):
        client, tmp_path = flask_client

        monkeypatch.setattr("web_app.TURNSTILE_ENABLED", True)
        monkeypatch.setattr("web_app.TURNSTILE_SECRET_KEY", "secret-key")
        mock_extract = MagicMock(return_value="should not run")
        monkeypatch.setattr("web_app.extract_text", mock_extract)

        data = {
            "file": (io.BytesIO(b"dummy"), "blocked.pdf"),
            "mode": "extract",
        }
        resp = client.post("/api/extract", data=data, content_type="multipart/form-data")

        assert resp.status_code == 400
        assert "Bot verification failed" in resp.get_json()["error"]
        mock_extract.assert_not_called()
        assert not (tmp_path / "uploads").exists()

    def test_extract_accepts_valid_turnstile_token(self, flask_client, monkeypatch):
        client, _ = flask_client

        monkeypatch.setattr("web_app.TURNSTILE_ENABLED", True)
        mock_verify = MagicMock(return_value={"success": True})
        monkeypatch.setattr("web_app.verify_turnstile_response", mock_verify)

        data = {
            "file": (io.BytesIO(b"dummy"), "verified.pdf"),
            "mode": "extract",
            "cf-turnstile-response": "client-token",
        }
        resp = client.post(
            "/api/extract",
            data=data,
            content_type="multipart/form-data",
            environ_overrides={"REMOTE_ADDR": "203.0.113.77"},
        )

        assert resp.status_code == 200
        mock_verify.assert_called_once_with("client-token", "203.0.113.77")

    def test_extract_rejects_pdf_over_page_limit_before_ocr(self, flask_client, monkeypatch):
        client, _ = flask_client
        import web_app

        mock_validate = MagicMock(
            side_effect=web_app.PdfPageLimitError(
                "PDFs are limited to 20 pages. This PDF has 21 pages."
            )
        )
        mock_extract = MagicMock(return_value="should not run")
        monkeypatch.setattr("web_app.validate_pdf_page_count", mock_validate)
        monkeypatch.setattr("web_app.extract_text", mock_extract)

        data = {
            "file": (io.BytesIO(b"dummy"), "too-many-pages.pdf"),
            "mode": "extract",
        }
        resp = client.post("/api/extract", data=data, content_type="multipart/form-data")

        assert resp.status_code == 400
        body = resp.get_json()
        assert "20 pages" in body["error"]
        mock_validate.assert_called_once()
        mock_extract.assert_not_called()

    def test_extract_png_success(self, flask_client):
        client, tmp_path = flask_client
        # Minimal 1×1 PNG bytes
        png_bytes = (
            b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01"
            b"\x00\x00\x00\x01\x08\x02\x00\x00\x00\x90wS\xde\x00\x00"
            b"\x00\x0cIDATx\x9cc\xf8\x0f\x00\x00\x01\x01\x00\x05\x18"
            b"\xd5N\x00\x00\x00\x00IEND\xaeB`\x82"
        )
        data = {
            "file": (io.BytesIO(png_bytes), "test.png"),
            "mode": "extract",
        }
        resp = client.post("/api/extract", data=data, content_type="multipart/form-data")
        assert resp.status_code == 200
        body = resp.get_json()
        assert body["text"] == "extracted text content"
        assert body["fileName"] == "test.png"
        assert "characterCount" in body
        assert "wordCount" in body

    def test_extract_does_not_persist_result(self, flask_client):
        client, tmp_path = flask_client
        data = {
            "file": (io.BytesIO(b"dummy"), "test.pdf"),
            "mode": "extract",
        }
        resp = client.post("/api/extract", data=data, content_type="multipart/form-data")
        assert resp.status_code == 200
        assert not (tmp_path / "output").exists()

    def test_extract_deletes_upload_after_processing(self, flask_client):
        client, tmp_path = flask_client
        data = {
            "file": (io.BytesIO(b"dummy"), "test.png"),
            "mode": "extract",
        }
        client.post("/api/extract", data=data, content_type="multipart/form-data")
        uploads = list((tmp_path / "uploads").iterdir()) if (tmp_path / "uploads").exists() else []
        assert uploads == [], "Uploaded file should be deleted after processing"

    def test_extract_returns_503_when_processing_slots_are_full(
        self, flask_client, monkeypatch
    ):
        client, tmp_path = flask_client
        import web_app

        mock_extract = MagicMock(return_value="should not run")
        monkeypatch.setattr("web_app.extract_text", mock_extract)

        assert web_app.acquire_processing_slot()
        try:
            data = {
                "file": (io.BytesIO(b"dummy"), "busy.pdf"),
                "mode": "extract",
            }
            resp = client.post("/api/extract", data=data, content_type="multipart/form-data")
        finally:
            web_app.release_processing_slot()

        assert resp.status_code == 503
        assert "Server is busy" in resp.get_json()["error"]
        mock_extract.assert_not_called()
        uploads = list((tmp_path / "uploads").iterdir()) if (tmp_path / "uploads").exists() else []
        assert uploads == []

    def test_extract_full_mode_calls_classify_and_summarize(self, flask_client, monkeypatch):
        client, tmp_path = flask_client

        mock_classify = MagicMock(return_value={"document_type": "invoice", "confidence": 90})
        mock_summarize = MagicMock(
            return_value={"short_summary": "A bill.", "bullet_points": ["Pay it."]}
        )
        monkeypatch.setattr("web_app.extract_text", MagicMock(return_value="invoice text here"))
        monkeypatch.setattr("agent.doc_classify.classify_document", mock_classify)
        monkeypatch.setattr("agent.doc_summarize.summarize_document", mock_summarize)

        data = {
            "file": (io.BytesIO(b"dummy"), "invoice.pdf"),
            "mode": "full",
        }
        resp = client.post("/api/extract", data=data, content_type="multipart/form-data")

        assert resp.status_code == 200
        body = resp.get_json()
        assert body["classification"] == {"document_type": "invoice", "confidence": 90}
        assert body["summary"]["short_summary"] == "A bill."
        assert body["summary"]["bullet_points"] == ["Pay it."]
        mock_classify.assert_called_once_with("invoice text here")
        mock_summarize.assert_called_once_with("invoice text here")

    def test_extract_full_mode_rate_limit_blocks_fourth_request(
        self, flask_client, monkeypatch
    ):
        client, _ = flask_client
        import web_app

        monkeypatch.setattr(web_app.limiter, "enabled", True)
        web_app.limiter.reset()
        mock_extract = MagicMock(return_value="invoice text here")
        monkeypatch.setattr("web_app.extract_text", mock_extract)
        monkeypatch.setattr(
            "agent.doc_classify.classify_document",
            MagicMock(return_value={"document_type": "invoice", "confidence": 90}),
        )
        monkeypatch.setattr(
            "agent.doc_summarize.summarize_document",
            MagicMock(return_value={"short_summary": "A bill.", "bullet_points": []}),
        )

        responses = []
        for _ in range(4):
            data = {
                "file": (io.BytesIO(b"dummy"), "invoice.pdf"),
                "mode": "full",
            }
            responses.append(
                client.post(
                    "/api/extract",
                    data=data,
                    content_type="multipart/form-data",
                    environ_overrides={"REMOTE_ADDR": "203.0.113.10"},
                )
            )

        assert [resp.status_code for resp in responses] == [200, 200, 200, 429]
        assert "Too many requests" in responses[-1].get_json()["error"]
        assert mock_extract.call_count == 3

    def test_extract_text_only_is_not_limited_by_full_mode_limit(
        self, flask_client, monkeypatch
    ):
        client, _ = flask_client
        import web_app

        monkeypatch.setattr(web_app.limiter, "enabled", True)
        web_app.limiter.reset()
        mock_extract = MagicMock(return_value="plain extracted text")
        monkeypatch.setattr("web_app.extract_text", mock_extract)

        responses = []
        for _ in range(4):
            data = {
                "file": (io.BytesIO(b"dummy"), "document.pdf"),
                "mode": "extract",
            }
            responses.append(
                client.post(
                    "/api/extract",
                    data=data,
                    content_type="multipart/form-data",
                    environ_overrides={"REMOTE_ADDR": "203.0.113.11"},
                )
            )

        assert [resp.status_code for resp in responses] == [200, 200, 200, 200]
        assert mock_extract.call_count == 4

    def test_extract_text_only_rate_limit_blocks_thirty_first_request(
        self, flask_client, monkeypatch
    ):
        client, _ = flask_client
        import web_app

        monkeypatch.setattr(web_app.limiter, "enabled", True)
        web_app.limiter.reset()
        mock_extract = MagicMock(return_value="plain extracted text")
        monkeypatch.setattr("web_app.extract_text", mock_extract)

        responses = []
        for _ in range(31):
            data = {
                "file": (io.BytesIO(b"dummy"), "document.pdf"),
                "mode": "extract",
            }
            responses.append(
                client.post(
                    "/api/extract",
                    data=data,
                    content_type="multipart/form-data",
                    environ_overrides={"REMOTE_ADDR": "203.0.113.13"},
                )
            )

        assert [resp.status_code for resp in responses[:30]] == [200] * 30
        assert responses[-1].status_code == 429
        assert "Too many requests" in responses[-1].get_json()["error"]
        assert mock_extract.call_count == 30

    def test_extract_full_mode_keeps_summary_when_classification_fails(
        self, flask_client, monkeypatch
    ):
        client, _ = flask_client

        mock_classify = MagicMock(side_effect=ValueError("bad classification json"))
        mock_summarize = MagicMock(
            return_value={"short_summary": "A bill.", "bullet_points": ["Pay it."]}
        )
        monkeypatch.setattr("web_app.extract_text", MagicMock(return_value="invoice text here"))
        monkeypatch.setattr("agent.doc_classify.classify_document", mock_classify)
        monkeypatch.setattr("agent.doc_summarize.summarize_document", mock_summarize)

        data = {
            "file": (io.BytesIO(b"dummy"), "invoice.pdf"),
            "mode": "full",
        }
        resp = client.post("/api/extract", data=data, content_type="multipart/form-data")

        assert resp.status_code == 200
        body = resp.get_json()
        assert body["text"] == "invoice text here"
        assert "classification" not in body
        assert body["summary"]["short_summary"] == "A bill."
        assert body["analysisErrors"] == {
            "classification": "Classification could not be completed."
        }
        assert "bad classification json" not in json.dumps(body)
        mock_classify.assert_called_once_with("invoice text here")
        mock_summarize.assert_called_once_with("invoice text here")

    def test_extract_full_mode_keeps_classification_when_summary_fails(
        self, flask_client, monkeypatch
    ):
        client, _ = flask_client

        mock_classify = MagicMock(return_value={"document_type": "invoice", "confidence": 90})
        mock_summarize = MagicMock(side_effect=ValueError("bad summary json"))
        monkeypatch.setattr("web_app.extract_text", MagicMock(return_value="invoice text here"))
        monkeypatch.setattr("agent.doc_classify.classify_document", mock_classify)
        monkeypatch.setattr("agent.doc_summarize.summarize_document", mock_summarize)

        data = {
            "file": (io.BytesIO(b"dummy"), "invoice.pdf"),
            "mode": "full",
        }
        resp = client.post("/api/extract", data=data, content_type="multipart/form-data")

        assert resp.status_code == 200
        body = resp.get_json()
        assert body["text"] == "invoice text here"
        assert body["classification"] == {"document_type": "invoice", "confidence": 90}
        assert "summary" not in body
        assert body["analysisErrors"] == {"summary": "Summary could not be completed."}
        assert "bad summary json" not in json.dumps(body)
        mock_classify.assert_called_once_with("invoice text here")
        mock_summarize.assert_called_once_with("invoice text here")

    def test_extract_failure_returns_500_json_and_deletes_upload(self, flask_client, monkeypatch):
        client, tmp_path = flask_client
        monkeypatch.setattr(
            "web_app.extract_text", MagicMock(side_effect=RuntimeError("ocr exploded"))
        )

        data = {
            "file": (io.BytesIO(b"dummy"), "test.pdf"),
            "mode": "extract",
        }
        resp = client.post("/api/extract", data=data, content_type="multipart/form-data")

        assert resp.status_code == 500
        body = resp.get_json()
        assert body["error"] == "Something went wrong while processing the document. Please try again."
        response_text = resp.get_data(as_text=True)
        assert "ocr exploded" not in response_text
        assert "RuntimeError" not in response_text
        uploads = list((tmp_path / "uploads").iterdir()) if (tmp_path / "uploads").exists() else []
        assert uploads == [], "Upload must be deleted even when processing fails"

    def test_extract_non_ascii_filename_succeeds(self, flask_client):
        client, _ = flask_client
        data = {
            "file": (io.BytesIO(b"dummy"), "файл.pdf"),
            "mode": "extract",
        }
        resp = client.post("/api/extract", data=data, content_type="multipart/form-data")

        assert resp.status_code == 200
        body = resp.get_json()
        assert body["text"] == "extracted text content"
        # secure_filename strips the Cyrillic; a sane fallback name is used
        assert body["fileName"].endswith(".pdf")

    def test_oversized_upload_returns_json_413(self, flask_client, monkeypatch):
        client, _ = flask_client
        import web_app

        monkeypatch.setitem(web_app.app.config, "MAX_CONTENT_LENGTH", 100)
        data = {
            "file": (io.BytesIO(b"x" * 1000), "big.pdf"),
            "mode": "extract",
        }
        resp = client.post("/api/extract", data=data, content_type="multipart/form-data")

        assert resp.status_code == 413
        body = resp.get_json()
        assert body is not None, "413 responses must be JSON so the frontend can show them"
        assert "error" in body

    def test_explain_no_file_returns_400(self, flask_client):
        client, _ = flask_client
        resp = client.post("/api/explain", data={"language": "ukrainian"})
        assert resp.status_code == 400
        assert "error" in resp.get_json()

    def test_explain_unsupported_type_returns_400(self, flask_client):
        client, _ = flask_client
        data = {
            "file": (io.BytesIO(b"data"), "file.xlsx"),
            "language": "ukrainian",
        }
        resp = client.post("/api/explain", data=data, content_type="multipart/form-data")
        assert resp.status_code == 400

    def test_explain_unsupported_language_returns_400(self, flask_client):
        client, _ = flask_client
        data = {
            "file": (io.BytesIO(b"dummy"), "test.pdf"),
            "language": "klingon",
        }
        resp = client.post("/api/explain", data=data, content_type="multipart/form-data")
        assert resp.status_code == 400

    def test_explain_can_be_disabled_before_upload_save(self, flask_client, monkeypatch):
        client, tmp_path = flask_client

        monkeypatch.setattr("web_app.OPENAI_FEATURES_ENABLED", False)
        mock_extract = MagicMock(return_value="should not run")
        monkeypatch.setattr("web_app.extract_text", mock_extract)
        mock_verify = MagicMock()
        monkeypatch.setattr("web_app.verify_request_turnstile", mock_verify)

        data = {
            "file": (io.BytesIO(b"dummy"), "blocked.pdf"),
            "language": "english",
        }
        resp = client.post("/api/explain", data=data, content_type="multipart/form-data")

        assert resp.status_code == 503
        assert "AI document features are temporarily unavailable" in resp.get_json()["error"]
        mock_verify.assert_not_called()
        mock_extract.assert_not_called()
        assert not (tmp_path / "uploads").exists()

    def test_explain_requires_turnstile_before_saving_upload(self, flask_client, monkeypatch):
        client, tmp_path = flask_client

        monkeypatch.setattr("web_app.TURNSTILE_ENABLED", True)
        monkeypatch.setattr("web_app.TURNSTILE_SECRET_KEY", "secret-key")
        mock_extract = MagicMock(return_value="should not run")
        monkeypatch.setattr("web_app.extract_text", mock_extract)

        data = {
            "file": (io.BytesIO(b"dummy"), "blocked.pdf"),
            "language": "english",
        }
        resp = client.post("/api/explain", data=data, content_type="multipart/form-data")

        assert resp.status_code == 400
        assert "Bot verification failed" in resp.get_json()["error"]
        mock_extract.assert_not_called()
        assert not (tmp_path / "uploads").exists()

    def test_explain_rejects_pdf_over_page_limit_before_ocr(self, flask_client, monkeypatch):
        client, _ = flask_client
        import web_app

        mock_validate = MagicMock(
            side_effect=web_app.PdfPageLimitError(
                "PDFs are limited to 20 pages. This PDF has 21 pages."
            )
        )
        mock_extract = MagicMock(return_value="should not run")
        monkeypatch.setattr("web_app.validate_pdf_page_count", mock_validate)
        monkeypatch.setattr("web_app.extract_text", mock_extract)

        data = {
            "file": (io.BytesIO(b"dummy"), "too-many-pages.pdf"),
            "language": "english",
        }
        resp = client.post("/api/explain", data=data, content_type="multipart/form-data")

        assert resp.status_code == 400
        body = resp.get_json()
        assert "20 pages" in body["error"]
        mock_validate.assert_called_once()
        mock_extract.assert_not_called()

    def test_explain_rate_limit_blocks_fourth_request(self, flask_client, monkeypatch):
        client, _ = flask_client
        import web_app

        monkeypatch.setattr(web_app.limiter, "enabled", True)
        web_app.limiter.reset()
        mock_extract = MagicMock(return_value="document text")
        monkeypatch.setattr("web_app.extract_text", mock_extract)
        monkeypatch.setattr(
            "agent.doc_explain.explain_document",
            MagicMock(
                return_value={
                    "document_type": "Letter",
                    "summary": "A letter.",
                    "explanation": "The sender shares information.",
                    "important_points": [],
                    "actions_required": [],
                    "important_dates": [],
                    "amounts": [],
                    "warnings": [],
                }
            ),
        )

        responses = []
        for _ in range(4):
            data = {
                "file": (io.BytesIO(b"dummy"), "letter.pdf"),
                "language": "english",
            }
            responses.append(
                client.post(
                    "/api/explain",
                    data=data,
                    content_type="multipart/form-data",
                    environ_overrides={"REMOTE_ADDR": "203.0.113.12"},
                )
            )

        assert [resp.status_code for resp in responses] == [200, 200, 200, 429]
        assert "Too many requests" in responses[-1].get_json()["error"]
        assert mock_extract.call_count == 3

    def test_explain_returns_503_when_processing_slots_are_full(
        self, flask_client, monkeypatch
    ):
        client, tmp_path = flask_client
        import web_app

        mock_extract = MagicMock(return_value="should not run")
        monkeypatch.setattr("web_app.extract_text", mock_extract)

        assert web_app.acquire_processing_slot()
        try:
            data = {
                "file": (io.BytesIO(b"dummy"), "busy.pdf"),
                "language": "english",
            }
            resp = client.post("/api/explain", data=data, content_type="multipart/form-data")
        finally:
            web_app.release_processing_slot()

        assert resp.status_code == 503
        assert "Server is busy" in resp.get_json()["error"]
        mock_extract.assert_not_called()
        uploads = list((tmp_path / "uploads").iterdir()) if (tmp_path / "uploads").exists() else []
        assert uploads == []

    def test_explain_empty_text_returns_422(self, flask_client, monkeypatch):
        client, _ = flask_client
        monkeypatch.setattr("web_app.extract_text", MagicMock(return_value="   \n  "))

        data = {
            "file": (io.BytesIO(b"dummy"), "blank.pdf"),
            "language": "english",
        }
        resp = client.post("/api/explain", data=data, content_type="multipart/form-data")

        assert resp.status_code == 422
        assert "error" in resp.get_json()

    def test_explain_truncates_long_documents(self, flask_client, monkeypatch):
        client, _ = flask_client
        import web_app

        long_text = "a" * (web_app.MAX_EXPLAIN_CHARS + 5000)
        monkeypatch.setattr("web_app.extract_text", MagicMock(return_value=long_text))

        explanation = {
            "document_type": "Contract",
            "summary": "A long contract.",
            "explanation": "It is long.",
            "important_points": [],
            "actions_required": [],
            "important_dates": [],
            "amounts": [],
            "warnings": [],
        }
        mock_explain = MagicMock(return_value=explanation)
        monkeypatch.setattr("agent.doc_explain.explain_document", mock_explain)

        data = {
            "file": (io.BytesIO(b"dummy"), "long.pdf"),
            "language": "english",
        }
        resp = client.post("/api/explain", data=data, content_type="multipart/form-data")

        assert resp.status_code == 200
        body = resp.get_json()
        assert body["sourceWasTruncated"] is True
        assert body["sourceCharacterCount"] == len(long_text)
        sent_text = mock_explain.call_args.args[0]
        assert len(sent_text) == web_app.MAX_EXPLAIN_CHARS

    def test_explain_failure_returns_500_json_and_deletes_upload(self, flask_client, monkeypatch):
        client, tmp_path = flask_client
        monkeypatch.setattr(
            "agent.doc_explain.explain_document",
            MagicMock(side_effect=RuntimeError("api exploded")),
        )

        data = {
            "file": (io.BytesIO(b"dummy"), "test.pdf"),
            "language": "english",
        }
        resp = client.post("/api/explain", data=data, content_type="multipart/form-data")

        assert resp.status_code == 500
        body = resp.get_json()
        assert body["error"] == "Something went wrong while processing the document. Please try again."
        response_text = resp.get_data(as_text=True)
        assert "api exploded" not in response_text
        assert "RuntimeError" not in response_text
        uploads = list((tmp_path / "uploads").iterdir()) if (tmp_path / "uploads").exists() else []
        assert uploads == [], "Upload must be deleted even when processing fails"

    def test_explain_english_skips_translation(self, flask_client, monkeypatch):
        client, _ = flask_client
        explanation = {
            "document_type": "Invoice",
            "summary": "This is an invoice.",
            "explanation": "The document asks for payment.",
            "important_points": ["Payment is due."],
            "actions_required": ["Pay the invoice."],
            "important_dates": [],
            "amounts": ["$500"],
            "warnings": [],
        }
        mock_explain = MagicMock(return_value=explanation)
        mock_translate = MagicMock()
        monkeypatch.setattr("agent.doc_explain.explain_document", mock_explain)
        monkeypatch.setattr("agent.doc_explain.translate_explanation", mock_translate)

        data = {
            "file": (io.BytesIO(b"dummy"), "test.pdf"),
            "language": "english",
        }
        resp = client.post("/api/explain", data=data, content_type="multipart/form-data")

        assert resp.status_code == 200
        body = resp.get_json()
        assert body["english"]["summary"] == "This is an invoice."
        assert body["translated"] is None
        mock_translate.assert_not_called()

    def test_explain_ukrainian_calls_translation(self, flask_client, monkeypatch):
        client, _ = flask_client
        explanation = {
            "document_type": "Invoice",
            "summary": "This is an invoice.",
            "explanation": "The document asks for payment.",
            "important_points": [],
            "actions_required": [],
            "important_dates": [],
            "amounts": [],
            "warnings": [],
        }
        translated = {
            **explanation,
            "document_type": "Рахунок",
            "summary": "Це рахунок.",
        }
        mock_translate = MagicMock(return_value=translated)
        monkeypatch.setattr("agent.doc_explain.explain_document", MagicMock(return_value=explanation))
        monkeypatch.setattr("agent.doc_explain.translate_explanation", mock_translate)

        data = {
            "file": (io.BytesIO(b"dummy"), "test.pdf"),
            "language": "ukrainian",
        }
        resp = client.post("/api/explain", data=data, content_type="multipart/form-data")

        assert resp.status_code == 200
        body = resp.get_json()
        assert body["languageName"] == "Ukrainian"
        assert body["translated"]["summary"] == "Це рахунок."
        mock_translate.assert_called_once_with(explanation, "Ukrainian")

    def test_explain_does_not_persist_result(self, flask_client, monkeypatch):
        client, tmp_path = flask_client
        explanation = {
            "document_type": "Letter",
            "summary": "This is a letter.",
            "explanation": "The sender is communicating information.",
            "important_points": [],
            "actions_required": [],
            "important_dates": [],
            "amounts": [],
            "warnings": [],
        }
        monkeypatch.setattr("agent.doc_explain.explain_document", MagicMock(return_value=explanation))

        data = {
            "file": (io.BytesIO(b"dummy"), "test.pdf"),
            "language": "english",
        }
        resp = client.post("/api/explain", data=data, content_type="multipart/form-data")
        assert resp.status_code == 200
        assert not (tmp_path / "output").exists()

    def test_explain_deletes_upload_after_processing(self, flask_client, monkeypatch, tmp_path):
        client, tmp_path = flask_client
        explanation = {
            "document_type": "Memo",
            "summary": "This is a memo.",
            "explanation": "The memo shares information.",
            "important_points": [],
            "actions_required": [],
            "important_dates": [],
            "amounts": [],
            "warnings": [],
        }
        monkeypatch.setattr("agent.doc_explain.explain_document", MagicMock(return_value=explanation))

        data = {
            "file": (io.BytesIO(b"dummy"), "test.pdf"),
            "language": "english",
        }
        client.post("/api/explain", data=data, content_type="multipart/form-data")

        uploads = list((tmp_path / "uploads").iterdir()) if (tmp_path / "uploads").exists() else []
        assert uploads == []


# ---------------------------------------------------------------------------
# 7. agent.py — process_single_file and main
# ---------------------------------------------------------------------------

class TestAgent:
    def test_process_single_file_extract_mode(self, monkeypatch, tmp_path):
        from agent.agent import process_single_file

        monkeypatch.setattr("agent.agent.extract_text", MagicMock(return_value="hello world text"))
        result = process_single_file("fake.pdf", "extract")

        assert result["file"] == "fake.pdf"
        assert result["extracted_text"] == "hello world text"
        assert "classification" not in result
        assert "summary" not in result

    def test_process_single_file_full_mode(self, monkeypatch):
        from agent.agent import process_single_file

        monkeypatch.setattr("agent.agent.extract_text", MagicMock(return_value="some text"))
        monkeypatch.setattr(
            "agent.agent.classify_document",
            MagicMock(return_value={"document_type": "other", "confidence": 70}),
        )
        monkeypatch.setattr(
            "agent.agent.summarize_document",
            MagicMock(return_value={"short_summary": "x", "bullet_points": []}),
        )

        result = process_single_file("fake.pdf", "full")

        assert result["classification"] == {"document_type": "other", "confidence": 70}
        assert result["summary"] == {"short_summary": "x", "bullet_points": []}

    def test_process_single_file_keeps_summary_when_classification_fails(self, monkeypatch):
        from agent.agent import process_single_file

        monkeypatch.setattr("agent.agent.extract_text", MagicMock(return_value="some text"))
        monkeypatch.setattr(
            "agent.agent.classify_document",
            MagicMock(side_effect=ValueError("bad classification json")),
        )
        monkeypatch.setattr(
            "agent.agent.summarize_document",
            MagicMock(return_value={"short_summary": "x", "bullet_points": []}),
        )

        result = process_single_file("fake.pdf", "full")

        assert result["extracted_text"] == "some text"
        assert "classification" not in result
        assert result["summary"] == {"short_summary": "x", "bullet_points": []}
        assert result["analysisErrors"] == {
            "classification": "Classification could not be completed."
        }

    def test_process_single_file_keeps_classification_when_summary_fails(self, monkeypatch):
        from agent.agent import process_single_file

        monkeypatch.setattr("agent.agent.extract_text", MagicMock(return_value="some text"))
        monkeypatch.setattr(
            "agent.agent.classify_document",
            MagicMock(return_value={"document_type": "other", "confidence": 70}),
        )
        monkeypatch.setattr(
            "agent.agent.summarize_document",
            MagicMock(side_effect=ValueError("bad summary json")),
        )

        result = process_single_file("fake.pdf", "full")

        assert result["extracted_text"] == "some text"
        assert result["classification"] == {"document_type": "other", "confidence": 70}
        assert "summary" not in result
        assert result["analysisErrors"] == {"summary": "Summary could not be completed."}

    def test_main_file_not_found(self, monkeypatch, capsys):
        import agent.agent as ag

        monkeypatch.setattr("sys.argv", ["agent", "nonexistent_path"])
        monkeypatch.setattr("os.path.isfile", MagicMock(return_value=False))
        monkeypatch.setattr("os.path.isdir", MagicMock(return_value=False))

        # Should print "Path not found" and return without raising
        with patch("builtins.open", mock_open()):
            ag.main()

        captured = capsys.readouterr()
        assert "Path not found" in captured.out

    def test_main_processes_directory(self, monkeypatch, tmp_path):
        import agent.agent as ag

        (tmp_path / "a.png").write_bytes(b"fake")
        (tmp_path / "b.png").write_bytes(b"fake")
        monkeypatch.setattr("sys.argv", ["agent", str(tmp_path)])
        monkeypatch.setattr("agent.agent.extract_text", MagicMock(return_value="text"))

        with patch("builtins.open", mock_open()):
            ag.main()

        assert ag.extract_text.call_count == 2

    def test_main_creates_output_directory(self, monkeypatch, tmp_path):
        """A fresh clone has no output/ folder; main() must create it."""
        import agent.agent as ag

        sample = tmp_path / "doc.png"
        sample.write_bytes(b"fake")
        monkeypatch.chdir(tmp_path)
        monkeypatch.setattr("sys.argv", ["agent", str(sample)])
        monkeypatch.setattr("agent.agent.extract_text", MagicMock(return_value="real text"))

        ag.main()

        results = json.loads((tmp_path / "output" / "results.json").read_text(encoding="utf-8"))
        assert results[0]["extracted_text"] == "real text"
        assert (tmp_path / "output" / "extracted_text.txt").exists()

    def test_main_skips_unsupported_files_in_directory(self, monkeypatch, tmp_path, capsys):
        import agent.agent as ag

        input_dir = tmp_path / "documents"
        input_dir.mkdir()
        (input_dir / "invoice.pdf").write_bytes(b"fake")
        (input_dir / "notes.txt").write_text("not supported", encoding="utf-8")
        monkeypatch.chdir(tmp_path)
        monkeypatch.setattr("sys.argv", ["agent", str(input_dir)])
        monkeypatch.setattr("agent.agent.extract_text", MagicMock(return_value="invoice text"))

        ag.main()

        results = json.loads((tmp_path / "output" / "results.json").read_text(encoding="utf-8"))
        captured = capsys.readouterr()
        assert len(results) == 1
        assert results[0]["file"].endswith("invoice.pdf")
        assert ag.extract_text.call_count == 1
        assert "Skipped notes.txt: unsupported file type" in captured.out
        assert "1 processed, 1 skipped, 0 failed" in captured.out

    def test_main_continues_after_file_failure(self, monkeypatch, tmp_path, capsys):
        import agent.agent as ag

        input_dir = tmp_path / "documents"
        input_dir.mkdir()
        (input_dir / "bad.pdf").write_bytes(b"bad")
        (input_dir / "good.pdf").write_bytes(b"good")
        monkeypatch.chdir(tmp_path)
        monkeypatch.setattr("sys.argv", ["agent", str(input_dir)])

        def extract_with_one_failure(path):
            if path.endswith("bad.pdf"):
                raise RuntimeError("corrupt document")
            return "good text"

        monkeypatch.setattr("agent.agent.extract_text", MagicMock(side_effect=extract_with_one_failure))

        ag.main()

        results = json.loads((tmp_path / "output" / "results.json").read_text(encoding="utf-8"))
        captured = capsys.readouterr()
        assert len(results) == 2
        assert results[0]["file"].endswith("bad.pdf")
        assert results[0]["error"] == "RuntimeError: corrupt document"
        assert "extracted_text" not in results[0]
        assert results[1]["file"].endswith("good.pdf")
        assert results[1]["extracted_text"] == "good text"
        assert ag.extract_text.call_count == 2
        assert "Failed bad.pdf: RuntimeError: corrupt document" in captured.out
        assert "1 processed, 0 skipped, 1 failed" in captured.out
        extracted_txt = (tmp_path / "output" / "extracted_text.txt").read_text(encoding="utf-8")
        assert "good.pdf" in extracted_txt
        assert "bad.pdf" not in extracted_txt


# ---------------------------------------------------------------------------
# 8. parse_jsonish (web_app utility)
# ---------------------------------------------------------------------------

class TestParseJsonish:
    def test_valid_json_string(self):
        import web_app
        result = web_app.parse_jsonish('{"key": "value"}')
        assert result == {"key": "value"}

    def test_invalid_json_returns_original_string(self):
        import web_app
        result = web_app.parse_jsonish("not json at all")
        assert result == "not json at all"

    def test_non_string_passthrough(self):
        import web_app
        assert web_app.parse_jsonish({"already": "dict"}) == {"already": "dict"}
        assert web_app.parse_jsonish(42) == 42
        assert web_app.parse_jsonish(None) is None


# ---------------------------------------------------------------------------
# 9. Integration tests (skipped unless --run-integration flag is passed)
# ---------------------------------------------------------------------------


@pytest.fixture()
def run_integration(request):
    if not request.config.getoption("--run-integration"):
        pytest.skip("Pass --run-integration to run this test")


class TestIntegration:
    """
    These tests hit real PaddleOCR and real sample images.
    Only run locally when you have the model downloaded.

        pytest tests/ --run-integration -v
    """

    def test_invoice_png_returns_nonempty_text(self, run_integration):
        from agent.text_extraction import extract_text
        path = str(SAMPLES_DIR / "invoice_test.png")
        result = extract_text(path)
        assert isinstance(result, str)
        assert len(result.strip()) > 0, "Expected non-empty OCR output from invoice_test.png"

    def test_ocr_test1_png_returns_nonempty_text(self, run_integration):
        from agent.text_extraction import extract_text
        path = str(SAMPLES_DIR / "ocr_test1.png")
        result = extract_text(path)
        assert len(result.strip()) > 0

    def test_unsupported_extension_raises_integration(self, run_integration):
        from agent.text_extraction import extract_text
        with pytest.raises(ValueError):
            extract_text("file.csv")
