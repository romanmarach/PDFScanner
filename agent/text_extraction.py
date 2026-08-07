import os
import tempfile

# Skip Paddle's model-source network probe during app startup.
os.environ.setdefault("PADDLE_PDX_DISABLE_MODEL_SOURCE_CHECK", "True")

import docx
import pdfplumber
from paddleocr import PaddleOCR


# Source of truth for what extract_text() can handle. The CLI and web app
# derive their accepted-file lists from this; the web app may expose a
# subset, but must never accept an extension that is not in this set.
EXTRACTABLE_EXTENSIONS = frozenset({".pdf", ".png", ".jpg", ".jpeg", ".docx"})
MIN_PAGE_TEXT_CHARS = 20


_ocr = None


def get_ocr():
    """Create the PaddleOCR model only when OCR is actually needed."""
    global _ocr

    if _ocr is None:
        # PaddleOCR 3.x init (cls/use_angle_cls is replaced by use_textline_orientation).
        _ocr = PaddleOCR(
            lang="en",
            use_textline_orientation=True,
            use_doc_orientation_classify=False,
            use_doc_unwarping=False,
        )

    return _ocr


def extract_text(file_path: str) -> str:
    """Extract text from a supported document path."""
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return extract_pdf(file_path)

    if ext in [".jpg", ".jpeg", ".png"]:
        return ocr_image(file_path)

    if ext == ".docx":
        return extract_docx(file_path)

    raise ValueError(f"Unsupported file type: {ext}")


def extract_pdf(path: str) -> str:
    """Extract each PDF page, using OCR only for pages with little embedded text."""
    with pdfplumber.open(path) as pdf:
        page_texts = [page.extract_text() or "" for page in pdf.pages]

    pages_needing_ocr = [
        page_num
        for page_num, text in enumerate(page_texts)
        if len(text.strip()) < MIN_PAGE_TEXT_CHARS
    ]

    if pages_needing_ocr:
        print(
            f"PDF has {len(pages_needing_ocr)} page(s) with little embedded text; "
            "using OCR for those pages..."
        )
        ocr_texts = ocr_pdf_pages(path, pages_needing_ocr)
        for page_num, ocr_text in ocr_texts.items():
            if ocr_text.strip():
                page_texts[page_num] = ocr_text

    return "\n\n".join(page_texts)


def paddle_predict_to_text(predict_output) -> str:
    """
    Convert PaddleOCR 3.x predict() output into plain text by extracting
    recognized text lines.
    """
    lines: list[str] = []

    for res in predict_output:
        j = getattr(res, "json", None)
        if not isinstance(j, dict):
            continue

        core = j.get("res", j)

        if isinstance(core, dict):
            rec_texts = core.get("rec_texts")
            if isinstance(rec_texts, list):
                lines.extend([text for text in rec_texts if isinstance(text, str)])
                continue

        if isinstance(core, list):
            for item in core:
                if isinstance(item, dict):
                    rec_texts = item.get("rec_texts")
                    if isinstance(rec_texts, list):
                        lines.extend([text for text in rec_texts if isinstance(text, str)])

    return "\n".join(lines)


def ocr_image(path: str) -> str:
    """OCR for image files using PaddleOCR 3.x."""
    print("Extracting from OCR image directly (PaddleOCR predict)")
    output = get_ocr().predict(input=path)
    return paddle_predict_to_text(output)


def ocr_pdf_pages(path: str, page_numbers) -> dict[int, str]:
    """Render and OCR selected zero-based PDF page numbers."""
    try:
        import fitz  # PyMuPDF
    except ImportError as exc:
        raise ImportError(
            "PyMuPDF is required for OCR fallback on PDFs. Install it with: pip install pymupdf"
        ) from exc

    requested_pages = set(page_numbers)
    page_texts: dict[int, str] = {}

    with fitz.open(path) as pdf:
        invalid_pages = requested_pages.difference(range(len(pdf)))
        if invalid_pages:
            raise ValueError(f"PDF page numbers out of range: {sorted(invalid_pages)}")

        with tempfile.TemporaryDirectory() as temp_dir:
            for page_num in sorted(requested_pages):
                page = pdf[page_num]
                pix = page.get_pixmap()

                temp_img = os.path.join(temp_dir, f"page_{page_num}.png")
                pix.save(temp_img)

                output = get_ocr().predict(input=temp_img)
                page_texts[page_num] = paddle_predict_to_text(output).strip()

    return page_texts


def ocr_pdf(path: str) -> str:
    """Render and OCR every page in a PDF."""
    try:
        import fitz  # PyMuPDF
    except ImportError as exc:
        raise ImportError(
            "PyMuPDF is required for OCR fallback on PDFs. Install it with: pip install pymupdf"
        ) from exc

    with fitz.open(path) as pdf:
        page_numbers = list(range(len(pdf)))

    page_texts = ocr_pdf_pages(path, page_numbers)
    # Technical debt: PDF pages are currently rendered at PyMuPDF's default
    # resolution, about 72 DPI, before OCR. Increasing the render scale with
    # fitz.Matrix(2, 2) or a configurable DPI may improve recognition of small,
    # faint, or low-quality text, but it also increases processing time and memory
    # usage. Current OCR accuracy is acceptable, so treat this as a future
    # accuracy/performance option rather than a bug.
    return "\n\n".join(page_texts[page_num] for page_num in page_numbers)


def extract_docx(path: str) -> str:
    """Extract text from Word documents, including tables, headers, and footers."""
    from docx.table import Table

    doc = docx.Document(path)
    blocks: list[str] = []

    for section in doc.sections:
        for para in section.header.paragraphs:
            if para.text.strip():
                blocks.append(para.text)

    for item in doc.iter_inner_content():
        if isinstance(item, Table):
            for row in item.rows:
                cells = [cell.text.strip() for cell in row.cells]
                blocks.append(" | ".join(cells))
        elif item.text.strip():
            blocks.append(item.text)

    for section in doc.sections:
        for para in section.footer.paragraphs:
            if para.text.strip():
                blocks.append(para.text)

    return "\n".join(blocks)
